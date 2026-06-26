"""
Document text extraction service
Supports: PDF (digital & scanned), DOCX, Images
"""

import io
from typing import Tuple


def extract_from_pdf(pdf_bytes: bytes) -> Tuple[str, bool]:
    """
    Extract text from PDF.
    Returns: (text, is_digital)
    - is_digital=True: PDF has selectable text
    - is_digital=False: PDF is scanned/image-based, needs OCR
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""

    for page in doc:
        page_text = page.get_text()
        text += page_text + "\n"

    doc.close()

    # If text is very short, likely scanned PDF
    is_digital = len(text.strip()) > 50

    return text.strip(), is_digital


def extract_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)


def get_file_type(filename: str) -> str:
    """Determine file type from filename"""
    filename_lower = filename.lower()

    if filename_lower.endswith('.pdf'):
        return 'pdf'
    elif filename_lower.endswith('.docx'):
        return 'docx'
    elif filename_lower.endswith('.doc'):
        return 'doc'
    elif filename_lower.endswith(('.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff')):
        return 'image'
    else:
        return 'unknown'


def extract_text(content: bytes, filename: str) -> Tuple[str, str, bool]:
    """
    Extract text from any supported file type.
    Returns: (text, file_type, needs_ocr)
    """
    file_type = get_file_type(filename)

    if file_type == 'pdf':
        text, is_digital = extract_from_pdf(content)
        if is_digital and len(text) > 50:
            return text, file_type, False
        else:
            # Scanned PDF, needs OCR
            return text, file_type, True

    elif file_type == 'docx':
        text = extract_from_docx(content)
        return text, file_type, False

    elif file_type == 'image':
        # Images always need OCR
        return "", file_type, True

    else:
        return "", file_type, True
